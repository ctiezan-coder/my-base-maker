#!/usr/bin/env python3
"""Script pour générer le guide d'utilisation ACIEX en format .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_guide():
    doc = Document()

    # Configuration des marges
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # === PAGE DE TITRE ===
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title_run = title.add_run("GUIDE D'UTILISATION")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(30, 91, 168)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Plateforme ACIEX")
    sub_run.bold = True
    sub_run.font.size = Pt(22)
    sub_run.font.color.rgb = RGBColor(255, 140, 0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc_run = desc.add_run("Côte d'Ivoire Export")
    desc_run.font.size = Pt(16)
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    version = doc.add_paragraph()
    version.add_run("Version 1.0 - Janvier 2026").font.size = Pt(11)
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # === TABLE DES MATIÈRES ===
    toc_title = doc.add_heading("Table des matières", level=1)

    toc_items = [
        ("1. Introduction", 3),
        ("2. Premiers pas", 3),
        ("3. Navigation", 4),
        ("4. Direction Générale", 5),
        ("5. DAF - Direction Administrative et Financière", 8),
        ("6. Directions Opérationnelles", 10),
        ("7. Administration", 12),
        ("8. Fonctionnalités communes", 13),
        ("9. Support", 14),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f"{item}").font.size = Pt(11)
        p.add_run(f"{'.' * (60 - len(item))} {page}").font.size = Pt(11)

    doc.add_page_break()

    # === 1. INTRODUCTION ===
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("Présentation", level=2)
    doc.add_paragraph(
        "ACIEX (Côte d'Ivoire Export) est une plateforme de gestion institutionnelle "
        "pour coordonner les activités de commerce et d'exportation, la gestion "
        "organisationnelle et le suivi opérationnel."
    )

    doc.add_heading("Fonctionnalités principales", level=2)
    features = [
        "Gestion des opérateurs économiques",
        "Suivi des projets et activités",
        "Gestion documentaire",
        "Messagerie interne",
        "Suivi budgétaire et comptable (SYSCOHADA)",
        "Gestion des ressources humaines",
        "Développement de marchés à l'export",
        "Tableaux de bord analytiques"
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # === 2. PREMIERS PAS ===
    doc.add_heading("2. Premiers pas", level=1)

    doc.add_heading("Création de compte", level=2)
    doc.add_paragraph(
        "1. Accédez à la page de connexion\n"
        "2. Cliquez sur \"S'inscrire\"\n"
        "3. Entrez votre email professionnel (@cotedivoirexport.ci)\n"
        "4. Créez un mot de passe sécurisé (min. 8 caractères avec majuscules, minuscules et chiffres)\n"
        "5. Validez l'inscription"
    )

    p = doc.add_paragraph()
    p.add_run("Important : ").bold = True
    p.add_run("Seules les adresses @cotedivoirexport.ci sont autorisées.")

    doc.add_heading("Approbation du compte", level=2)
    doc.add_paragraph(
        "Après inscription, un administrateur doit approuver votre compte. "
        "Statuts possibles : En attente, Approuvé, Rejeté."
    )

    doc.add_heading("Connexion", level=2)
    doc.add_paragraph("Entrez votre email et mot de passe sur la page de connexion.")

    doc.add_page_break()

    # === 3. NAVIGATION ===
    doc.add_heading("3. Navigation", level=1)

    doc.add_heading("Organisation de la barre latérale", level=2)
    doc.add_paragraph(
        "La navigation est organisée en 4 sections principales, chacune avec un code couleur :"
    )

    # Tableau des sections
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = table.rows[0].cells
    headers[0].text = "Section"
    headers[1].text = "Modules"

    sections_data = [
        ("Direction Générale (Orange)", "Tableau de bord, Messagerie, Opérateurs, Projets, Documents, KPIs"),
        ("DAF (Vert)", "Achats, Support, RH, Missions, Budgets, Comptabilité"),
        ("Directions Opérationnelles (Cyan)", "Marchés, Partenariats, Formations, Agenda, Événements"),
        ("Administration (Rouge)", "Gestion utilisateurs, Permissions"),
    ]

    for i, (section, modules) in enumerate(sections_data, 1):
        row = table.rows[i].cells
        row[0].text = section
        row[1].text = modules

    doc.add_paragraph()
    doc.add_heading("En-tête", level=2)
    doc.add_paragraph("Logo ACIEX, menu utilisateur (profil, déconnexion), notifications.")

    doc.add_page_break()

    # === 4. DIRECTION GÉNÉRALE ===
    doc.add_heading("4. Direction Générale", level=1)

    doc.add_heading("Tableau de bord", level=2)
    doc.add_paragraph(
        "Vue d'ensemble avec KPIs (entreprises, projets, opportunités), "
        "graphiques analytiques et accès rapide aux modules."
    )

    doc.add_heading("Messagerie", level=2)
    doc.add_paragraph(
        "Communication interne : conversations privées, historique, notifications de nouveaux messages."
    )

    doc.add_heading("Opérateurs Économiques", level=2)
    doc.add_paragraph("Gestion de la base de données des entreprises accompagnées.")
    actions = [
        "Recherche et filtres (nom, secteur, région)",
        "Import en masse via Excel",
        "Export des données",
        "Ajout/modification d'entreprises"
    ]
    for a in actions:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_heading("Projets", level=2)
    doc.add_paragraph(
        "Suivi des projets institutionnels avec statuts : En préparation, En cours, Terminé, Suspendu."
    )

    doc.add_heading("Documents", level=2)
    doc.add_paragraph(
        "Système de gestion documentaire avec dossiers hiérarchiques, versioning et métadonnées."
    )

    doc.add_heading("Suivi & Évaluation", level=2)
    doc.add_paragraph("Performance par direction : taux de réalisation, activités réalisées vs planifiées.")

    doc.add_heading("KPIs", level=2)
    doc.add_paragraph("Indicateurs clés de performance avec historique et alertes.")

    doc.add_heading("Archive des Activités", level=2)
    doc.add_paragraph("Journal chronologique de toutes les activités avec filtrage et export.")

    doc.add_page_break()

    # === 5. DAF ===
    doc.add_heading("5. DAF - Direction Administrative et Financière", level=1)

    doc.add_heading("Achats", level=2)
    doc.add_paragraph(
        "Gestion des fournisseurs et bons de commande. "
        "Statuts : Brouillon, En attente, Approuvé, Livré, Annulé."
    )

    doc.add_heading("Support & Maintenance", level=2)
    doc.add_paragraph(
        "Système de tickets avec priorités : Urgente, Haute, Moyenne, Basse. "
        "Suivi et historique des résolutions."
    )

    doc.add_heading("Ressources Humaines", level=2)
    doc.add_paragraph("Répertoire du personnel et gestion des congés.")
    rh_features = [
        "Liste des employés par direction",
        "Demandes de congés avec workflow d'approbation",
        "Suivi des présences"
    ]
    for f in rh_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading("Ordres de Missions", level=2)
    doc.add_paragraph(
        "Gestion des déplacements : objet, destination, dates, transport, budget. "
        "Statuts : En attente, Validé, En cours, Terminé."
    )

    doc.add_heading("Budgets", level=2)
    doc.add_paragraph(
        "Suivi budgétaire par direction : allocation, consommation, alertes de dépassement, prévisions."
    )

    doc.add_heading("Comptabilité", level=2)
    doc.add_paragraph("Système conforme au plan SYSCOHADA.")
    compta_features = [
        "Plan comptable avec classes et catégories",
        "Écritures comptables (débit/crédit)",
        "Journaux : achats, ventes, banque, OD"
    ]
    for f in compta_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # === 6. DIRECTIONS OPÉRATIONNELLES ===
    doc.add_heading("6. Directions Opérationnelles", level=1)

    doc.add_heading("Développement des Marchés", level=2)
    doc.add_paragraph("Gestion des opportunités d'export et connexions d'affaires.")
    market_features = [
        "Opportunités d'export internationales",
        "Connexions B2B et suivi des négociations",
        "Analyse des marchés potentiels avec cartographie"
    ]
    for f in market_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading("Partenariats", level=2)
    doc.add_paragraph(
        "Gestion des partenariats stratégiques : conventions, accords-cadres, protocoles. "
        "Suivi des échéances et engagements."
    )

    doc.add_heading("Formations", level=2)
    doc.add_paragraph(
        "Programmes de formation : thème, formateurs, dates, participants, évaluations d'impact."
    )

    doc.add_heading("Agenda", level=2)
    doc.add_paragraph("Calendrier institutionnel avec vues mensuelle, hebdomadaire, journalière.")

    doc.add_heading("Événements", level=2)
    doc.add_paragraph(
        "Gestion des événements : séminaires, conférences, salons, missions commerciales. "
        "Capacité, urgence, inscriptions."
    )

    doc.add_heading("Médiathèque", level=2)
    doc.add_paragraph("Galerie de contenus : photos, vidéos, présentations. Organisation par albums et tags.")

    doc.add_heading("Imputations", level=2)
    doc.add_paragraph(
        "Affectation des tâches avec responsables et deadlines. "
        "Statuts : À faire, En cours, Terminé, En retard."
    )

    doc.add_heading("Espace Collaborateurs", level=2)
    doc.add_paragraph("Espace collaboratif : discussions, partage de ressources, idées, annonces.")

    doc.add_page_break()

    # === 7. ADMINISTRATION ===
    doc.add_heading("7. Administration", level=1)

    p = doc.add_paragraph()
    p.add_run("Note : ").bold = True
    p.add_run("Modules réservés aux administrateurs.")

    doc.add_heading("Gestion des Utilisateurs", level=2)
    admin_features = [
        "Liste et recherche des utilisateurs",
        "Approbation/rejet des nouveaux comptes",
        "Gestion des emails autorisés",
        "Suspension/réactivation de comptes"
    ]
    for f in admin_features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading("Permissions", level=2)
    doc.add_paragraph("Gestion des rôles et droits d'accès par module.")

    # Tableau des rôles
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    headers = table.rows[0].cells
    headers[0].text = "Rôle"
    headers[1].text = "Accès"

    roles_data = [
        ("Administrateur", "Accès complet à tous les modules"),
        ("Manager", "Accès aux modules de sa direction"),
        ("Utilisateur", "Accès limité selon les permissions"),
    ]

    for i, (role, access) in enumerate(roles_data, 1):
        row = table.rows[i].cells
        row[0].text = role
        row[1].text = access

    doc.add_page_break()

    # === 8. FONCTIONNALITÉS COMMUNES ===
    doc.add_heading("8. Fonctionnalités communes", level=1)

    doc.add_heading("Recherche et filtres", level=2)
    doc.add_paragraph(
        "Barre de recherche par mots-clés, filtres avancés (statut, date, direction), "
        "tri par colonnes."
    )

    doc.add_heading("Export des données", level=2)
    doc.add_paragraph(
        "1. Appliquez vos filtres\n"
        "2. Cliquez sur \"Exporter\"\n"
        "3. Téléchargez le fichier Excel"
    )

    doc.add_heading("Pagination", level=2)
    doc.add_paragraph("Listes paginées (15 éléments par défaut). Navigation par flèches ou numéro de page.")

    doc.add_heading("Notifications", level=2)
    doc.add_paragraph("Alertes pour : nouveaux messages, approbations requises, échéances, mises à jour.")

    doc.add_heading("Mode sombre", level=2)
    doc.add_paragraph("Activez le mode sombre dans les préférences du profil.")

    doc.add_heading("Raccourcis clavier", level=2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    shortcuts = [
        ("Action", "Raccourci"),
        ("Recherche rapide", "Ctrl + K"),
        ("Nouveau message", "Ctrl + M"),
        ("Tableau de bord", "Ctrl + H"),
    ]

    for i, (action, shortcut) in enumerate(shortcuts):
        row = table.rows[i].cells
        row[0].text = action
        row[1].text = shortcut
        if i == 0:
            row[0].paragraphs[0].runs[0].bold = True
            row[1].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # === 9. SUPPORT ===
    doc.add_heading("9. Support", level=1)

    doc.add_heading("FAQ", level=2)

    faqs = [
        ("Mot de passe oublié ?", "Cliquez sur \"Mot de passe oublié\" sur la page de connexion."),
        ("Compte en attente ?", "Contactez l'administrateur de votre direction."),
        ("Accès manquant à un module ?", "Les accès dépendent de votre rôle. Contactez l'administrateur."),
        ("Import d'entreprises ?", "Utilisez l'import en masse dans \"Opérateurs économiques\"."),
    ]

    for q, a in faqs:
        p = doc.add_paragraph()
        p.add_run(f"Q: {q}").bold = True
        doc.add_paragraph(f"R: {a}")

    doc.add_heading("Contact Support", level=2)
    doc.add_paragraph(
        "Pour toute assistance :\n"
        "1. Accédez au module \"Support & Maintenance\"\n"
        "2. Créez un ticket avec description détaillée\n"
        "3. Joignez des captures d'écran si nécessaire"
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("© ACIEX - Côte d'Ivoire Export - Tous droits réservés")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    # Sauvegarde
    doc.save('GUIDE_UTILISATEUR_ACIEX.docx')
    print("Guide généré avec succès : GUIDE_UTILISATEUR_ACIEX.docx")

if __name__ == "__main__":
    create_guide()
